import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\Users\dell\Desktop\NEW APP")
RESULTS = ROOT / "test-artifacts" / "e2e-results.json"
OUTPUT = ROOT / "test-artifacts" / "End to End Testing.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D9E1E8"
TEXT = "243447"
MUTED = "5F6B7A"
GREEN = "16794B"
GREEN_FILL = "E8F5EE"
RED = "9B1C1C"
RED_FILL = "FDECEC"
AMBER = "7A5A00"
AMBER_FILL = "FFF5D6"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="CBD5E1", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=11, color=TEXT, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_text(cell, text, *, bold=False, color=TEXT, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    set_run_font(p.add_run(str(text)), size=size, color=color, bold=bold)


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run_font(p.add_run(text))
    return p


def add_status_callout(doc, label, message, status):
    fill, color = (GREEN_FILL, GREEN) if status == "PASS" else (RED_FILL, RED) if status == "FAIL" else (AMBER_FILL, AMBER)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size="3")
    shade(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{label}: "), size=11, color=color, bold=True)
    set_run_font(p.add_run(message), size=10.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def status_color(status):
    return GREEN if status == "PASS" else RED if status == "FAIL" else AMBER


def status_fill(status):
    return GREEN_FILL if status == "PASS" else RED_FILL if status == "FAIL" else AMBER_FILL


data = json.loads(RESULTS.read_text(encoding="utf-8"))
api_tests = data["tests"]

manual_tests = [
    {"id":"BLD-001","module":"Backend Build","name":"NestJS production compilation","status":"PASS","details":"npm run build completed successfully."},
    {"id":"BLD-002","module":"Admin Build","name":"Next.js production compilation and TypeScript","status":"PASS","details":"Production build and static generation completed successfully."},
    {"id":"BLD-003","module":"Mobile Static","name":"Mobile TypeScript validation","status":"PASS","details":"npx tsc --noEmit completed with zero errors."},
    {"id":"BLD-004","module":"Mobile Static","name":"Expo ESLint validation","status":"PASS","details":"npx expo lint completed with zero errors."},
    {"id":"BLD-005","module":"Admin Static","name":"Admin ESLint validation","status":"PASS","details":"npm run lint completed with zero errors."},
    {"id":"BLD-006","module":"Android Build","name":"Android debug APK compilation","status":"PASS","details":"Gradle assembleDebug completed successfully; Razorpay and Expo notifications native modules compiled."},
    {"id":"BLD-007","module":"Backend Unit Tests","name":"Repository Jest suite","status":"BLOCKED","details":"No *.spec.ts files exist; Jest returned 'No tests found'. API regression coverage was supplied separately."},
    {"id":"UI-001","module":"Admin UI","name":"Administrator authentication and dashboard load","status":"PASS","details":"Temporary isolated test administrator authenticated; dashboard rendered without UI alerts."},
    {"id":"UI-002","module":"Admin UI","name":"Enquiry Support list and filters","status":"PASS","details":"Ticket list, status filter, category filter and attachment counts rendered."},
    {"id":"UI-003","module":"Admin UI","name":"Support ticket multiple-photo gallery","status":"PASS","details":"A three-photo fixture displayed USER ATTACHMENTS (3) with three individually clickable images."},
    {"id":"UI-004","module":"Admin UI","name":"QR Generator quick statistics","status":"PASS","details":"Total 11,710; scanned 522; active 11,188 rendered from API data."},
    {"id":"UI-005","module":"Admin UI","name":"QR Generator product search dropdown","status":"PASS","details":"Clicking search loaded the active product list with SKU, category and points."},
    {"id":"UI-006","module":"Admin UI","name":"Product management page","status":"PASS","details":"Products page and product search control rendered with no UI alert."},
    {"id":"UI-007","module":"Admin UI","name":"Delivery Tracker workflow screen","status":"PASS","details":"Pending, Payment Done, Dispatched, Delivered and Rejected states rendered."},
    {"id":"UI-008","module":"Admin UI","name":"Notification management screen","status":"PASS","details":"Notification page and audience targeting controls rendered without UI alerts."},
    {"id":"UI-009","module":"Admin UI","name":"App Settings push notification controls","status":"PASS","details":"Push notification controls rendered without UI alerts."},
    {"id":"UI-010","module":"Admin UI","name":"Electrician bonus and type filters","status":"PASS","details":"Welcome Back Bonus, All Electrician and App Status filters rendered."},
    {"id":"UI-011","module":"Admin UI","name":"Electrician API data population","status":"PASS","details":"Loading completed with 50 visible table rows and no load error."},
    {"id":"MOB-001","module":"Mobile Support","name":"Visible plus-button for multiple photos","status":"PASS","details":"Need Help shows an explicit + control, count and multiple-photo guidance."},
    {"id":"MOB-002","module":"Mobile Support","name":"Multi-select gallery configuration","status":"PASS","details":"Image picker enables multiple selection and limits remaining capacity."},
    {"id":"MOB-003","module":"Mobile Support","name":"Five-photo maximum","status":"PASS","details":"Client and server both enforce a maximum of five attachments."},
    {"id":"MOB-004","module":"Mobile Support","name":"Selected-photo preview and removal","status":"PASS","details":"All selected photos are previewed; each attached thumbnail has a remove control."},
    {"id":"MOB-005","module":"Mobile Support","name":"Ticket detail attachment rendering","status":"PASS","details":"Submitted ticket detail displays all retained attachment URLs."},
    {"id":"MOB-006","module":"Admin Support","name":"Backward-compatible photo ingestion","status":"PASS","details":"Admin merges legacy photoUrl and new photoUrls without duplicates."},
    {"id":"PHY-001","module":"Physical Device","name":"Gallery multi-select on a real Android handset","status":"BLOCKED","details":"No Android handset was connected during this run; source, API and native build checks passed."},
    {"id":"PHY-002","module":"Physical Device","name":"OS push notification delivery","status":"BLOCKED","details":"No registered connected device token was available for physical receipt verification."},
    {"id":"PHY-003","module":"Payment","name":"Complete Razorpay UPI payment on handset","status":"BLOCKED","details":"Native Razorpay module compiled, but external UPI handoff requires a connected handset and Razorpay test transaction."},
]

all_tests = api_tests + manual_tests
passed = sum(1 for test in all_tests if test["status"] == "PASS")
failed = sum(1 for test in all_tests if test["status"] == "FAIL")
blocked = sum(1 for test in all_tests if test["status"] == "BLOCKED")

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for level, size, color, before, after in ((1,16,BLUE,16,8),(2,13,BLUE,12,6),(3,12,DARK_BLUE,8,4)):
    style = styles[f"Heading {level}"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run_font(hp.add_run("SRV ELECTRICALS  |  QUALITY ASSURANCE"), size=8.5, color=MUTED, bold=True)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(fp.add_run("End-to-End Test Report  |  Page "), size=8.5, color=MUTED)
add_field(fp, "PAGE")

# Memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(2)
set_run_font(p.add_run("END-TO-END TEST REPORT"), size=25, color=NAVY, bold=True)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
set_run_font(p.add_run("Mobile Application, Admin Panel and Backend API"), size=14, color=MUTED)

metadata = [
    ("Project", "SRV Electricals Digital Platform"),
    ("Test date", "18 June 2026"),
    ("Environment", "Local development services connected to the project database"),
    ("Scope", "Mobile app, admin UI, backend API, database integrity and Android build"),
    ("Prepared for", "Management and deployment review"),
    ("Overall verdict", "CONDITIONAL GO - functional suite passes; inventory data requires remediation"),
]
meta_table = doc.add_table(rows=len(metadata), cols=2)
set_table_geometry(meta_table, [1800, 7560])
set_table_borders(meta_table, color=MID_GRAY, size="3")
for row, (label, value) in zip(meta_table.rows, metadata):
    shade(row.cells[0], LIGHT_GRAY)
    set_cell_text(row.cells[0], label, bold=True, color=NAVY, size=9.5)
    set_cell_text(row.cells[1], value, size=9.5)

doc.add_paragraph()
metric_table = doc.add_table(rows=2, cols=4)
set_table_geometry(metric_table, [2340,2340,2340,2340], indent=0)
set_table_borders(metric_table, color=WHITE, size="2")
metrics = [("129", "TESTS"), (str(passed), "PASSED"), (str(failed), "FAILED"), (str(blocked), "BLOCKED")]
for idx, (value, label) in enumerate(metrics):
    color = NAVY if idx == 0 else GREEN if idx == 1 else RED if idx == 2 else AMBER
    fill = LIGHT_BLUE if idx == 0 else GREEN_FILL if idx == 1 else RED_FILL if idx == 2 else AMBER_FILL
    shade(metric_table.cell(0, idx), fill)
    shade(metric_table.cell(1, idx), fill)
    set_cell_text(metric_table.cell(0, idx), value, bold=True, color=color, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(metric_table.cell(1, idx), label, bold=True, color=color, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

add_heading(doc, "1. Executive Summary", 1)
add_body(doc, "The tested platform is functionally stable across the mobile application, administrator interface and backend API. All builds and static checks passed. The automated regression suite completed 102 API and data checks with 101 passes and one genuine failure. Eleven administrator UI checks and six focused multiple-photo checks also passed.")
add_status_callout(doc, "Release recommendation", "Proceed to controlled device acceptance testing after inventory stock values are corrected. Do not treat the current database as production-ready for product ordering.", "BLOCKED")
add_status_callout(doc, "Implemented feature", "Need Help now supports up to five photos with a visible plus button, multi-selection, previews, removal and ticket-history display. Enquiry Support receives and displays every image in a lazy-loaded gallery.", "PASS")

add_heading(doc, "2. Scope and Test Approach", 1)
for item in [
    "Backend contract and authorization checks against the running NestJS API.",
    "Role-based mobile API checks for electrician, dealer, customer and counter boy identities.",
    "Data-integrity checks for products, banners, QR inventory and product orders.",
    "Browser-based administrator regression for support, QR, products, delivery, notifications, settings and electrician filters.",
    "Production/static compilation checks for backend, admin, mobile TypeScript, lint and Android native code.",
    "Focused realistic multi-photo payload test, database verification and administrator gallery verification.",
]:
    add_bullet(doc, item)

add_heading(doc, "3. Feature Verification - Multiple Support Photos", 1)
feature_table = doc.add_table(rows=1, cols=3)
set_table_geometry(feature_table, [2400,5160,1800])
set_table_borders(feature_table)
headers = ["Layer", "Verified behaviour", "Result"]
for idx, header_text in enumerate(headers):
    shade(feature_table.cell(0, idx), NAVY)
    set_cell_text(feature_table.cell(0, idx), header_text, bold=True, color=WHITE, size=9)
set_repeat_table_header(feature_table.rows[0])
feature_rows = [
    ("Mobile UI", "Visible + button, 0/5 counter, multiple selection, all-photo confirmation preview and individual removal.", "PASS"),
    ("Mobile transport", "Sends photoUrls array and legacy photoUrl for backward compatibility.", "PASS"),
    ("Backend", "Accepts and stores up to five unique photo values in a PostgreSQL text array.", "PASS"),
    ("Realistic API payload", "Three approximately 260 KB images submitted successfully with HTTP 201.", "PASS"),
    ("Database", "cardinality(photoUrls) verified as 3 before test fixture cleanup.", "PASS"),
    ("Admin API", "Ticket detail returned all three photo values.", "PASS"),
    ("Admin UI", "Attachment badge displayed 3 photos; detail drawer rendered three clickable images.", "PASS"),
]
for layer, behaviour, result in feature_rows:
    cells = feature_table.add_row().cells
    set_cell_text(cells[0], layer, bold=True, size=9)
    set_cell_text(cells[1], behaviour, size=9)
    shade(cells[2], status_fill(result))
    set_cell_text(cells[2], result, bold=True, color=status_color(result), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
set_table_geometry(feature_table, [2400,5160,1800])

add_heading(doc, "4. Results by Area", 1)
area_rows = [
    ("Backend API", "101 / 102 passed", "One inventory data failure; all exercised endpoints responded correctly."),
    ("Mobile role APIs", "48 / 48 passed", "Authentication and core role services passed for all four roles."),
    ("Admin safe GET APIs", "30 / 30 passed", "All parameter-free administrator read endpoints in the OpenAPI contract passed."),
    ("Admin browser regression", "11 / 11 passed", "Targeted screens loaded with live API data and no visible error alerts."),
    ("Build and static analysis", "6 / 6 passed", "Backend, admin, mobile lint/type checks and Android build passed."),
    ("Physical acceptance", "0 / 3 executed", "Requires a connected handset and external payment/push services."),
]
area_table = doc.add_table(rows=1, cols=3)
set_table_geometry(area_table, [2300,1800,5260])
set_table_borders(area_table)
for idx, text in enumerate(("Area", "Outcome", "Assessment")):
    shade(area_table.cell(0, idx), NAVY)
    set_cell_text(area_table.cell(0, idx), text, bold=True, color=WHITE, size=9)
set_repeat_table_header(area_table.rows[0])
for area, outcome, assessment in area_rows:
    cells = area_table.add_row().cells
    set_cell_text(cells[0], area, bold=True, size=9)
    set_cell_text(cells[1], outcome, bold=True, color=GREEN if "passed" in outcome else AMBER, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cells[2], assessment, size=9)
set_table_geometry(area_table, [2300,1800,5260])

add_heading(doc, "5. Defect and Risk Register", 1)
risk_table = doc.add_table(rows=1, cols=5)
set_table_geometry(risk_table, [900,1050,1900,3510,2000])
set_table_borders(risk_table)
for idx, text in enumerate(("ID", "Severity", "Area", "Finding / impact", "Recommended action")):
    shade(risk_table.cell(0, idx), NAVY)
    set_cell_text(risk_table.cell(0, idx), text, bold=True, color=WHITE, size=8.5)
set_repeat_table_header(risk_table.rows[0])
risks = [
    ("BUG-001", "HIGH", "Inventory", "272 of 273 active products have stock <= 0. Ordering can incorrectly report products as unavailable.", "Load correct stock or agree and implement a non-stock-controlled ordering policy before deployment."),
    ("QA-001", "MEDIUM", "Backend tests", "Jest has no unit/integration test files. Regression protection currently depends on external smoke testing.", "Add service, controller and payment/support integration tests to CI."),
    ("DATA-001", "MEDIUM", "Categories", "Dedicated product_categories table contains 0 rows although the mobile API derives 30 categories from products.", "Confirm intended source of truth and populate/synchronize the category table before production administration."),
    ("ENV-001", "MEDIUM", "Push", "No physical device token was available, so OS-level notification receipt was not executed.", "Install the current native build, log in, confirm token registration, then send role and individual test notifications."),
    ("ENV-002", "MEDIUM", "Payment", "Razorpay native compilation passed, but UPI handoff and callback require physical-device acceptance.", "Run success, cancel and failure flows using Razorpay test mode on a connected handset."),
    ("TECH-001", "LOW", "Android build", "Gradle reports deprecated features that will be incompatible with Gradle 9.0.", "Record dependency upgrades in the next maintenance sprint; no current build failure."),
]
for rid, severity, area, finding, action in risks:
    cells = risk_table.add_row().cells
    set_cell_text(cells[0], rid, bold=True, size=8)
    sev_color = RED if severity == "HIGH" else AMBER if severity == "MEDIUM" else MUTED
    shade(cells[1], RED_FILL if severity == "HIGH" else AMBER_FILL if severity == "MEDIUM" else LIGHT_GRAY)
    set_cell_text(cells[1], severity, bold=True, color=sev_color, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cells[2], area, bold=True, size=8)
    set_cell_text(cells[3], finding, size=8)
    set_cell_text(cells[4], action, size=8)
set_table_geometry(risk_table, [900,1050,1900,3510,2000])

add_heading(doc, "6. Deployment Acceptance Checklist", 1)
checklist = [
    ("PASS", "Backend and admin production builds complete without compilation errors."),
    ("PASS", "Mobile TypeScript, Expo lint, admin lint and native Android compilation pass."),
    ("PASS", "Multiple support photos persist through app -> API -> database -> admin gallery."),
    ("ACTION", "Correct active product stock data and retest order creation."),
    ("ACTION", "Perform connected-device push, image picker and Razorpay acceptance tests."),
    ("ACTION", "Decide and document the production source of truth for product categories."),
]
for status, item in checklist:
    color = GREEN if status == "PASS" else AMBER
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run(f"{status}: "), color=color, bold=True)
    set_run_font(p.add_run(item))

add_heading(doc, "Appendix A - Complete Test Case Register", 1)
add_body(doc, "This appendix lists every executed, failed and blocked test included in the reported totals. Temporary test tickets and the temporary test administrator were removed after verification.")

test_table = doc.add_table(rows=1, cols=5)
widths = [720, 1500, 3420, 900, 2820]
set_table_geometry(test_table, widths)
set_table_borders(test_table, color="D5DCE3", size="3")
for idx, text in enumerate(("ID", "Module", "Test case", "Result", "Evidence")):
    shade(test_table.cell(0, idx), NAVY)
    set_cell_text(test_table.cell(0, idx), text, bold=True, color=WHITE, size=8)
set_repeat_table_header(test_table.rows[0])
for test in all_tests:
    cells = test_table.add_row().cells
    set_cell_text(cells[0], test["id"], bold=True, size=7.6, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cells[1], test["module"], size=7.6)
    set_cell_text(cells[2], test["name"], size=7.6)
    shade(cells[3], status_fill(test["status"]))
    set_cell_text(cells[3], test["status"], bold=True, color=status_color(test["status"]), size=7.6, align=WD_ALIGN_PARAGRAPH.CENTER)
    evidence = test.get("details", "")
    if test.get("durationMs") is not None:
        evidence += f" ({test['durationMs']} ms)"
    set_cell_text(cells[4], evidence, size=7.4)
set_table_geometry(test_table, widths)

add_heading(doc, "Appendix B - Test Data Snapshot", 1)
snapshot = data["summary"]["catalog"]
snapshot_rows = [
    ("Products", str(snapshot["productsTotal"]), f"{snapshot['productsActive']} active; {snapshot['activeZeroStock']} active with zero stock"),
    ("Derived product categories", "30", "Returned by mobile product-category API"),
    ("Dedicated category records", str(snapshot["categories"]), "Database product_categories table"),
    ("Active banners", str(snapshot["activeBanners"]), "All active banner API checks passed"),
    ("QR codes", f"{snapshot['qrTotal']:,}", f"{snapshot['qrScanned']:,} scanned; {snapshot['qrTotal']-snapshot['qrScanned']:,} remaining"),
]
snap_table = doc.add_table(rows=1, cols=3)
set_table_geometry(snap_table, [2400,1560,5400])
set_table_borders(snap_table)
for idx, text in enumerate(("Dataset", "Count", "Observation")):
    shade(snap_table.cell(0, idx), NAVY)
    set_cell_text(snap_table.cell(0, idx), text, bold=True, color=WHITE, size=9)
set_repeat_table_header(snap_table.rows[0])
for dataset, count, observation in snapshot_rows:
    cells = snap_table.add_row().cells
    set_cell_text(cells[0], dataset, bold=True, size=9)
    set_cell_text(cells[1], count, bold=True, color=NAVY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(cells[2], observation, size=9)
set_table_geometry(snap_table, [2400,1560,5400])

add_heading(doc, "Sign-off Statement", 1)
add_body(doc, "The evidence supports a conditional deployment decision. Application functionality, API integration and the newly requested multiple-photo support passed the available automated and browser checks. Production release should follow only after the inventory data issue and the three physical-device acceptance items are closed or formally accepted by the release owner.")

doc.core_properties.title = "End to End Testing"
doc.core_properties.subject = "SRV Electricals mobile, admin and backend quality assurance report"
doc.core_properties.author = "SRV Electricals QA"
doc.core_properties.keywords = "end-to-end testing, mobile, admin, backend, API, deployment"
doc.save(OUTPUT)
print(OUTPUT)
